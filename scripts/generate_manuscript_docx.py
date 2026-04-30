from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_manuscript() -> Document:
    doc = Document()

    title = doc.add_paragraph(
        "A Multi-Model Reliability Screening and Consensus-Aggregation Decision Framework for Large Language Model Evaluation"
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    subtitle = doc.add_paragraph(
        "Methodology Draft for SSCI Submission (Library and Information Science)"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Abstract", level=1)
    add_paragraph(
        doc,
        "This study proposes and validates a four-layer decision-governance framework for multi-evaluator large language model (LLM) assessment in library and information science (LIS). The framework is designed as a governance tool rather than an optimal classifier. Layer 1 screens evaluator stability with ICC(2,1) and confidence-interval-assisted boundary handling. Layer 2 enforces group-level agreement using Kendall's W with a greedy disagreement-removal procedure and MDS diagnostics. Layer 3 integrates median score and majority-consensus dispersion (rank IQR) in a nine-quadrant action matrix under a 20/40/40 percentile policy. Layer 4 provides optional skill-level diagnosis for disputed items. In a stress-test simulation of 10,000 books, 100 evaluators, and two independent rounds, 90 evaluators passed reliability screening and 85 remained after consensus aggregation (W=0.801). Replacing rank range with rank IQR shifted direct-accept recommendations from near-collapse to 1,887 books (18.9%). Bootstrap analysis (1,000 iterations) showed highly stable partitions, with the default 20/40/40 scheme at mean stability 0.9858 and only approximately 0.003 below the best searched configuration. Ten-seed Monte Carlo tests indicated low variability (final W mean 0.8065, SD 0.0031). We discuss validity risks, including reliability-validity separation, minority-school exclusion under greedy pruning, relative quantile boundaries, method-viewpoint divergence, and the need for external validation on real datasets.",
    )

    add_heading(doc, "Keywords", level=1)
    add_paragraph(
        doc,
        "LLM evaluation; decision governance; ICC; Kendall's W; consensus aggregation; rank IQR; collection development; library and information science",
    )

    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(
        doc,
        "Multi-model LLM evaluation often conflates score averaging with evaluator quality control. In practical LIS procurement and adoption scenarios, evaluators may be unstable, partially adversarial, or systematically biased. A governance-oriented method must therefore separate (i) evaluator reliability, (ii) evaluator agreement, and (iii) decision actionability. Existing studies provide useful statistics for reliability and agreement but do not always organize them into an end-to-end decision pipeline with explicit action outputs.",
    )
    add_paragraph(
        doc,
        "This paper contributes a four-layer framework to bridge this gap. The design objective is not to maximize predictive accuracy against a latent truth; instead, it is to produce accountable, robust, and interpretable decisions under evaluator plurality.",
    )

    add_heading(doc, "2. Theoretical Position and Framework Design", level=1)
    add_heading(doc, "2.1 Governance-Oriented Framing", level=2)
    add_paragraph(
        doc,
        "We position the framework as a decision governance tool. The core question is: which books are endorsed by sufficiently many stable evaluators under controlled agreement constraints? This shifts emphasis from classifier optimality to process transparency, reproducibility, and policy interpretability.",
    )

    add_heading(doc, "2.2 Four-Layer Architecture", level=2)
    add_numbered(
        doc,
        [
            "Layer 1 (Reliability): retain evaluators with ICC(2,1) >= 0.80, while using confidence intervals for boundary handling.",
            "Layer 2 (Consensus): require Kendall's W >= 0.80; if violated, iteratively remove the most disagreement-prone evaluator until threshold satisfaction or infeasible subset size.",
            "Layer 3 (Decision): map each book into a nine-quadrant matrix using median score and rank IQR (Q3-Q1), with 20/40/40 percentile partitioning and action labels.",
            "Layer 4 (Optional diagnosis): decompose disputed items by LIS skill dimensions (scholarly value, topical relevance, readability, authority/credibility, collection fit).",
        ],
    )

    add_heading(doc, "2.3 Key Methodological Decisions", level=2)
    add_bullets(
        doc,
        [
            "Consensus dispersion uses rank IQR instead of rank range, re-defining high consensus as majority consistency rather than unanimity.",
            "Default agreement threshold is theta_con = 0.80, with sensitivity scanning from 0.55 to 0.90.",
            "Default partition policy is 20/40/40; bootstrap evidence supports it as a stable governance default.",
        ],
    )

    add_heading(doc, "3. Simulation Stress-Test Design", level=1)
    add_heading(doc, "3.1 Data-Generation Scale and Structure", level=2)
    add_paragraph(
        doc,
        "The stress test includes 10,000 books and 100 evaluators over two independent rounds. Evaluators are grouped into ten archetypes (A-J) to emulate practical LIS evaluation risks: reliability gradients, adversarial disruptors, plural schools, missingness patterns, heteroscedastic noise, temporal drift, redundancy, and mixed extremes.",
    )

    add_heading(doc, "3.2 Evaluated Outputs", level=2)
    add_bullets(
        doc,
        [
            "Retained evaluator counts and confidence diagnostics in Layers 1-2.",
            "Nine-quadrant distributions and direct-action counts in Layer 3.",
            "Bootstrap partition stability across percentile schemes.",
            "Threshold sensitivity trajectories under shared greedy-path replay.",
            "Monte Carlo variability across 10 seeds.",
            "Alternative partition strategy comparisons (bootstrap-optimal percentile, GMM, consensus-depth voting).",
        ],
    )

    add_heading(doc, "4. Results", level=1)
    add_heading(doc, "4.1 Core Pipeline Results", level=2)
    add_bullets(
        doc,
        [
            "Layer 1 retained 90 evaluators.",
            "Layer 2 retained 85 evaluators at theta_con = 0.80 with Kendall's W = 0.801.",
            "Layer 3 produced 1,887 high-score/high-consensus books (18.9%), with zero data-insufficient books in this run.",
            "Low-score/high-consensus books reached 1,440, supporting explicit negative recommendations.",
        ],
    )

    add_heading(doc, "4.2 Stability and Robustness", level=2)
    add_paragraph(
        doc,
        "Bootstrap (1,000 iterations) showed high and close stability across schemes: 15/35/50 = 0.9864, 25/50/25 = 0.9863, 20/40/40 = 0.9859. The default policy remains in a stable region and is approximately 0.003 below the searched best level, indicating practical robustness with interpretability continuity.",
    )
    add_paragraph(
        doc,
        "Threshold scan (0.55-0.90) showed a controlled trade-off. At theta=0.80, retained evaluators were 85 and high-score/high-consensus count remained strong (1,886 in sensitivity replay).",
    )

    add_heading(doc, "4.3 Monte Carlo (10 Seeds)", level=2)
    add_paragraph(
        doc,
        "Across seeds [42, 52, 62, 72, 82, 92, 102, 112, 122, 132], metric volatility was low: reliable evaluators mean 90.0 (SD 0.0), consensus evaluators mean 85.1 (SD 0.3), final W mean 0.8065 (SD 0.0031), direct-accept count mean 1,894.7 (SD 17.66), and clear-reject count mean 1,395.7 (SD 29.00).",
    )

    add_heading(doc, "4.4 Alternative Layer-3 Methods", level=2)
    add_paragraph(
        doc,
        "Compared with baseline percentile (20/40/40, rank IQR), alternative methods produced distinct recommendation profiles:",
    )
    add_bullets(
        doc,
        [
            "Bootstrap-optimal percentile (p_high=0.15, p_low=0.35): high-high = 1,446; Jaccard with baseline = 0.7663.",
            "GMM cluster mapping: high-high = 1,316; Jaccard = 0.6974.",
            "Consensus-depth voting (>=70% evaluators place item in top-50%): high-high = 2,000; Jaccard = 0.9435.",
        ],
    )
    add_paragraph(
        doc,
        "These findings support the baseline as a balanced default while justifying complementary diagnostic views.",
    )

    add_heading(doc, "5. Discussion", level=1)
    add_heading(doc, "5.1 Why IQR Matters for Majority Consensus", level=2)
    add_paragraph(
        doc,
        "Under evaluator plurality, rank range over-penalizes minority extremes and can collapse actionable recommendations. Rank IQR filters tails and captures central agreement mass, making it better aligned with governance objectives in collection decisions.",
    )

    add_heading(
        doc, "5.2 Interpretation of Density vs Percentile Disagreement", level=2
    )
    add_paragraph(
        doc,
        "Low overlap between density-based and percentile-based direct-accept sets should not be interpreted as model failure. The density method reflects local manifold concentration, while the percentile method imposes global governance proportions. They serve different decision viewpoints.",
    )

    add_heading(doc, "5.3 Stated Validity Risks", level=2)
    add_numbered(
        doc,
        [
            "Reliability is not validity: stable evaluators can still be consistently wrong.",
            "Greedy pruning can exclude coherent minority schools.",
            "Quantile thresholds are relative boundaries and may drift under distribution shift.",
            "Density vs percentile divergence is viewpoint-driven, not necessarily quality-driven.",
            "Synthetic stress tests require external validation on real LIS corpora and workflows.",
        ],
    )

    add_heading(doc, "6. Conclusion", level=1)
    add_paragraph(
        doc,
        "This study demonstrates that an explicitly layered governance pipeline can convert multi-model LLM scoring into operationally actionable recommendations while preserving methodological transparency. The reliability-consensus-decision-diagnosis decomposition supports auditability and policy communication. Empirically, the majority-consensus upgrade via rank IQR resolves recommendation collapse and remains robust under bootstrap, threshold scans, alternative partition tests, and 10-seed Monte Carlo stress tests. For deployment, we recommend retaining percentile+IQR as default policy, with periodic bootstrap-based retuning as governance maintenance.",
    )

    add_heading(doc, "References", level=1)
    add_bullets(
        doc,
        [
            "Koo, T. K., & Li, M. Y. (2016). A Guideline of Selecting and Reporting Intraclass Correlation Coefficients for Reliability Research. Journal of Chiropractic Medicine, 15(2), 155-163.",
            "Landis, J. R., & Koch, G. G. (1977). The Measurement of Observer Agreement for Categorical Data. Biometrics, 33(1), 159-174.",
            "Shrout, P. E., & Fleiss, J. L. (1979). Intraclass Correlations: Uses in Assessing Rater Reliability. Psychological Bulletin, 86(2), 420-428.",
            "Thelwall, M., & Cox, J. (2025). Evaluating large language model outputs in scholarly information contexts: methodological considerations. (Referenced as theoretical context in this draft).",
        ],
    )

    add_heading(doc, "Appendix A. Reproducibility Snapshot", level=1)
    add_bullets(
        doc,
        [
            "Project: llm_eval_simulation",
            "Core reports: results/reports/summary_report.txt, monte_carlo_report.txt, alternative_methods_report.txt",
            "Main pipeline command: python main.py",
            'Alternative-method benchmark command: python -c "from src.alternative_methods import run_alternative_methods_report; run_alternative_methods_report()"',
        ],
    )

    return doc


def main() -> None:
    project_root = Path(__file__).resolve().parents[1]
    out_dir = project_root / "results" / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "manuscript_draft.docx"

    doc = build_manuscript()
    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()
